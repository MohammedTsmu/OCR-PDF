from docx import Document  # Word document creation
from docx.shared import Pt  # For setting font size
from docx.oxml.ns import qn  # For setting font name
from docx.shared import RGBColor  # For setting font color
from reportlab.pdfgen import canvas  # PDF creation
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import inch
from ocr.image_processing import pdf_to_images, image_to_text

def convert_to_text(pdf_path):
    images = pdf_to_images(pdf_path)
    text = '\n'.join([image_to_text(image) for image in images])
    return text

def convert_to_txt(pdf_path, result_path):
    text = convert_to_text(pdf_path)
    with open(result_path, 'w', encoding='utf-8') as file:
        file.write(text)

def convert_to_word(pdf_path, result_path):
    text = convert_to_text(pdf_path)
    doc = Document()
    
    # Set default font and size
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    font.color.rgb = RGBColor(0, 0, 0)
    
    # Add the extracted text with paragraphs
    for line in text.split('\n'):
        paragraph = doc.add_paragraph()
        run = paragraph.add_run(line)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
    
    doc.save(result_path)

def convert_to_pdf(pdf_path, result_path):
    text = convert_to_text(pdf_path)
    
    c = canvas.Canvas(result_path, pagesize=A4)
    width, height = A4

    # Set up margins
    x_margin = 1 * inch
    y_margin = 1 * inch
    text_width = width - 2 * x_margin
    y_position = height - y_margin

    # Split text into lines to handle wrapping
    lines = text.split('\n')
    for line in lines:
        wrapped_lines = c.beginText(x_margin, y_position)
        wrapped_lines.setFont("Helvetica", 12)
        wrapped_lines.setTextOrigin(x_margin, y_position)

        # Word wrap long lines
        for text_line in line.splitlines():
            while len(text_line) > 0:
                line_width = c.stringWidth(text_line, "Helvetica", 12)
                if line_width < text_width:
                    wrapped_lines.textLine(text_line)
                    text_line = ""
                else:
                    wrap_point = len(text_line)
                    while wrap_point > 0 and c.stringWidth(text_line[:wrap_point], "Helvetica", 12) > text_width:
                        wrap_point -= 1
                    wrapped_lines.textLine(text_line[:wrap_point])
                    text_line = text_line[wrap_point:]
            y_position -= 14  # Move the y position for the next line

        c.drawText(wrapped_lines)
        y_position -= 14  # Move the y position for the next paragraph

        # Check if the page is full, and create a new page if necessary
        if y_position < y_margin:
            c.showPage()
            y_position = height - y_margin

    c.save()

def convert_to_txt(pdf_path, result_path):
    text = convert_to_text(pdf_path)
    with open(result_path, 'w', encoding='utf-8') as file:
        file.write(text)
